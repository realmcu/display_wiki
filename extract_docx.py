#!/usr/bin/env python3
from docx import Document
import sys

def extract_to_markdown(docx_path, md_path):
    doc = Document(docx_path)
    
    with open(md_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                f.write('\n')
                continue
            
            # 处理标题
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                if level.isdigit():
                    f.write('#' * int(level) + ' ' + text + '\n\n')
                else:
                    f.write('# ' + text + '\n\n')
            else:
                f.write(text + '\n\n')
        
        # 提取表格
        if doc.tables:
            f.write('\n---\n\n## 表格内容\n\n')
            for i, table in enumerate(doc.tables, 1):
                f.write(f'### 表格 {i}\n\n')
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    f.write('| ' + ' | '.join(cells) + ' |\n')
                    if row == table.rows[0]:
                        f.write('| ' + ' | '.join(['---'] * len(cells)) + ' |\n')
                f.write('\n')

if __name__ == '__main__':
    docx_file = 'Bee4 Display controller SPEC_20221125.docx'
    md_file = 'Bee4_Display_Controller_SPEC.md'
    
    print(f'正在提取 {docx_file}...')
    extract_to_markdown(docx_file, md_file)
    print(f'已生成 {md_file}')
